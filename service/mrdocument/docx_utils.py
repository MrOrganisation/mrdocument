"""DOCX utilities for text extraction and metadata embedding."""

import io
import logging
from typing import Optional

from docx import Document
from docx.opc.exceptions import PackageNotFoundError

logger = logging.getLogger(__name__)


def extract_text_from_docx(docx_bytes: bytes) -> str:
    """
    Extract plain text from a DOCX file.

    Args:
        docx_bytes: Raw bytes of the DOCX file

    Returns:
        Extracted text content

    Raises:
        ValueError: If DOCX cannot be read
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except PackageNotFoundError as e:
        raise ValueError(f"Invalid DOCX file: {e}")
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {e}")

    # Extract text from paragraphs
    paragraphs = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            paragraphs.append(text)

    # Extract text from tables
    for table in doc.tables:
        for row in table.rows:
            row_text = []
            for cell in row.cells:
                cell_text = cell.text.strip()
                if cell_text:
                    row_text.append(cell_text)
            if row_text:
                paragraphs.append(" | ".join(row_text))

    text = "\n\n".join(paragraphs)
    logger.debug(f"Extracted {len(text)} chars from DOCX ({len(paragraphs)} paragraphs)")
    return text


def embed_metadata_in_docx(
    docx_bytes: bytes,
    title: Optional[str] = None,
    author: Optional[str] = None,
    subject: Optional[str] = None,
    keywords: Optional[list[str]] = None,
) -> bytes:
    """
    Embed metadata into a DOCX file's core properties.

    Args:
        docx_bytes: Raw bytes of the DOCX file
        title: Document title
        author: Document author/sender
        subject: Document subject
        keywords: List of keywords

    Returns:
        Modified DOCX file as bytes

    Raises:
        ValueError: If DOCX cannot be modified
    """
    try:
        doc = Document(io.BytesIO(docx_bytes))
    except Exception as e:
        raise ValueError(f"Failed to read DOCX: {e}")

    # Access core properties
    core_props = doc.core_properties

    if title:
        core_props.title = title
    if author:
        core_props.author = author
    if subject:
        core_props.subject = subject
    if keywords:
        core_props.keywords = ", ".join(keywords)

    # Save to bytes
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    
    result = output.getvalue()
    logger.debug(f"Embedded metadata in DOCX: title={title}, author={author}")
    return result
